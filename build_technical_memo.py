#!/usr/bin/env python3
"""
Build ClauseChain_Technical_Memo.docx for the UN Global Hackathon submission.

Placeholder PNGs are generated only when not already present, preserving real
screenshots and diagrams. Target: < 750 words visible in the document.
"""

from pathlib import Path
from PIL import Image, ImageDraw, ImageFont

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
ROOT = Path(__file__).parent
FIG_DIR = ROOT / "technical_memo_figures"
FIG_DIR.mkdir(exist_ok=True)
OUT_DOCX = ROOT / "ClauseChain_Technical_Memo.docx"

# ---------------------------------------------------------------------------
# Figures  (filename, docx caption, placeholder overlay text)
# ---------------------------------------------------------------------------
FIGURES = [
    ("fig01_dashboard.png",
     "Figure 1 — Workspace dashboard: coverage, KPIs, and pipeline health.",
     "Workspace dashboard\n/dashboard"),
    ("fig02_pipeline_stepper.png",
     "Figure 2 — Nine-stage pipeline stepper (Discover → Export) shown across all pipeline pages.",
     "Pipeline stepper\ndiscover -> export"),
    ("fig03_crawl_console.png",
     "Figure 3 — Crawl Console: candidate documents with type tags and discovery confidence.",
     "Crawl Console\n/pipeline/crawl"),
    ("fig04_harvest_review.png",
     "Figure 4 — Harvest Review: human keep/discard triage before extraction begins.",
     "Harvest Review\n/pipeline/harvest"),
    ("fig05_extraction_ocr_diff.png",
     "Figure 5 — Extraction Workspace: dual-engine OCR diff (Qwen2-VL vs. Tesseract) on a scanned Bengali page.",
     "Extraction Workspace\n/pipeline/extract  -  OCR diff"),
    ("fig06_mapping_run.png",
     "Figure 6 — Mapping Run: autonomy selector (L0–L3) and CVR gates streaming in real time.",
     "Mapping Run\n/pipeline/map  -  CVR loop"),
    ("fig07_source_trace.png",
     "Figure 7 — Source Trace: dual-panel sync between extracted clauses and original source, with per-gate results.",
     "Source Trace\n/pipeline/trace  -  dual-panel sync"),
    ("fig08_export_output.png",
     "Figure 8 — Export Output: JSONL / CSV records with discovery tags, CVR gate results, and verbatim citations.",
     "Export Output\n/pipeline/export"),
]

# Diagrams from build_diagrams.py — never overwritten as placeholders.
DIAGRAM_ARCH = ("fig00_architecture.png",
                "Figure 0 — System architecture: pipeline stages and CVR loop with local/cloud routing.")
DIAGRAM_CVR  = ("fig09_cvr_loop.png",
                "Figure 9 — CVR loop: every claim passes G1 Span, G2 Entailment, and G3 Structure before export.")

# ---------------------------------------------------------------------------
# 1. Generate placeholder PNGs only when missing (preserves real screenshots)
# ---------------------------------------------------------------------------
def make_placeholder(path: Path, overlay_text: str) -> None:
    W, H = 1600, 900
    img = Image.new("RGB", (W, H), color=(240, 242, 247))
    draw = ImageDraw.Draw(img)

    draw.rectangle([0, 0, W - 1, H - 1], outline=(180, 186, 196), width=4)
    draw.rectangle([24, 24, W - 25, H - 25], outline=(210, 215, 222), width=2)

    chip_w, chip_h = 220, 56
    draw.rectangle([40, 40, 40 + chip_w, 40 + chip_h], fill=(37, 99, 235))
    try:
        chip_font = ImageFont.truetype(
            "/System/Library/Fonts/Supplemental/Arial Bold.ttf", 22)
    except OSError:
        chip_font = ImageFont.load_default()
    draw.text((58, 54), "ClauseChain", fill=(255, 255, 255), font=chip_font)

    try:
        title_font = ImageFont.truetype(
            "/System/Library/Fonts/Supplemental/Arial Bold.ttf", 64)
        sub_font = ImageFont.truetype(
            "/System/Library/Fonts/Supplemental/Arial.ttf", 32)
    except OSError:
        title_font = ImageFont.load_default()
        sub_font = ImageFont.load_default()

    lines = overlay_text.split("\n")
    title_line = lines[0]
    sub_line = lines[1] if len(lines) > 1 else ""

    bbox_t = draw.textbbox((0, 0), title_line, font=title_font)
    tw, th = bbox_t[2] - bbox_t[0], bbox_t[3] - bbox_t[1]
    bbox_s = draw.textbbox((0, 0), sub_line, font=sub_font)
    sw, _  = bbox_s[2] - bbox_s[0], bbox_s[3] - bbox_s[1]

    total_h = th + 20 + (bbox_s[3] - bbox_s[1])
    cx, cy = W // 2, H // 2

    draw.text((cx - tw // 2, cy - total_h // 2),
              title_line, fill=(31, 41, 55), font=title_font)
    draw.text((cx - sw // 2, cy - total_h // 2 + th + 20),
              sub_line, fill=(107, 114, 128), font=sub_font)

    try:
        hint_font = ImageFont.truetype(
            "/System/Library/Fonts/Supplemental/Arial.ttf", 22)
    except OSError:
        hint_font = ImageFont.load_default()
    hint = "Placeholder — replace with screenshot (right-click → Change Picture in Word)"
    bbox_h = draw.textbbox((0, 0), hint, font=hint_font)
    draw.text((cx - (bbox_h[2] - bbox_h[0]) // 2, H - 70),
              hint, fill=(140, 146, 156), font=hint_font)

    img.save(path, "PNG", optimize=True)


print("Checking placeholder figures...")
for filename, _caption, overlay in FIGURES:
    p = FIG_DIR / filename
    if not p.exists():
        make_placeholder(p, overlay)
        print(f"  {filename} (generated placeholder)")
    else:
        print(f"  {filename} (kept existing)")

# ---------------------------------------------------------------------------
# 2. Build the .docx
# ---------------------------------------------------------------------------
doc = Document()

for section in doc.sections:
    section.top_margin    = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin   = Inches(0.9)
    section.right_margin  = Inches(0.9)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

for level, size in [(1, 16), (2, 12)]:
    s = doc.styles[f"Heading {level}"]
    s.font.name = "Calibri"
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)


def add_caption(text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x55, 0x5B, 0x66)
    p.paragraph_format.space_after = Pt(10)


def add_figure(filename: str, caption: str, width_in: float = 6.2) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(FIG_DIR / filename), width=Inches(width_in))
    p.paragraph_format.space_after = Pt(2)
    add_caption(caption)


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, h in enumerate(headers):
        hdr_cells = table.rows[0].cells
        hdr_cells[i].text = ""
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            cell.paragraphs[0].add_run(val).font.size = Pt(9)
    doc.add_paragraph()


# ---- TITLE BLOCK ---------------------------------------------------------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("ClauseChain — Technical Memo")
tr.bold = True
tr.font.size = Pt(20)
tr.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run(
    "Hash-Anchored Regulatory Evidence for Digital Trade  |  UN Global Hackathon: "
    "AI for Digital Trade Regulatory Analysis  |  RDTII Pillars 6 & 7"
)
sr.italic = True
sr.font.size = Pt(10)
sr.font.color.rgb = RGBColor(0x55, 0x5B, 0x66)

# ---- LEAD PARAGRAPH ------------------------------------------------------
doc.add_paragraph(
    "ClauseChain is an open-source AI pipeline that discovers, extracts, and maps "
    "digital trade regulations to the UN RDTII framework. Hallucination resistance "
    "is structural, not prompted: every classification is bound to a verbatim, "
    "hash-anchored citation from an authoritative source and passes a three-gate "
    "Cite-Verify-Reject (CVR) loop. A failed gate rejects rather than softens."
)

add_figure("fig01_dashboard.png", FIGURES[0][1])

# ---- §1 SYSTEM ARCHITECTURE ---------------------------------------------
doc.add_heading("1. System Architecture", level=1)

add_figure(DIAGRAM_ARCH[0], DIAGRAM_ARCH[1], width_in=6.4)

doc.add_paragraph(
    "ClauseChain implements six stages: Collect, Extract, Classify, Explain, Cite, "
    "Export. A seed registry drives Crawl4AI across official gazettes and statute "
    "databases. A file-type router dispatches HTML to a legal-structure parser, "
    "native PDFs to Docling, and scanned PDFs to dual-engine OCR (Qwen2-VL + "
    "Tesseract, consensus), emitting canonical JSON with character offsets and "
    "bounding boxes. Clauses embed into Qdrant via BGE-M3; Llama 3.1 fills a "
    "constrained JSON schema — no free prose, no paraphrase."
)

add_table(
    ["Stage", "Output"],
    [
        ("Discovery",        "crawled documents, source confidence-scored"),
        ("Extraction",       "canonical JSON with hierarchy, char offsets, bbox, SHA-256"),
        ("Indexing",         "BGE-M3 embeddings → Qdrant hybrid index"),
        ("Mapping",          "constrained schema: pillar, verbatim_span, rule, exceptions, confidence"),
        ("CVR Verification", "three-gate pass/reject with gate scores logged"),
        ("Export",           "JSONL / CSV / provenance bundle for independent re-verification"),
    ],
)

doc.add_paragraph(
    "The nine-page interface — dashboard, jurisdiction detail, audit view, RDTII "
    "matrix, pipeline ledger, crawl console, harvest triage, extraction workspace, "
    "mapping run — is a working Next.js / React prototype."
)

add_figure("fig02_pipeline_stepper.png", FIGURES[1][1])

# ---- §2 TOOLS, MODELS, AND METHODS --------------------------------------
doc.add_heading("2. Tools, Models, and Methods", level=1)

doc.add_paragraph(
    "ClauseChain defaults to self-hosted open-weight models on a single 24 GB GPU, "
    "with optional per-task cloud routing to OpenAI or Anthropic, capped by a "
    "per-run token budget."
)

add_table(
    ["Task", "Local default", "Cloud (optional)"],
    [
        ("Crawling",           "Crawl4AI",                             "n/a"),
        ("HTML parsing",       "markdown + legal-structure parser",    "n/a"),
        ("PDF extraction",     "Docling",                              "n/a"),
        ("OCR (scanned PDF)",  "Qwen2-VL 7B + Tesseract (consensus)", "n/a"),
        ("Embeddings",         "BGE-M3",                               "text-embedding-3-large"),
        ("LLM classification", "Llama 3.1 8B",                        "gpt-4.1 / claude-sonnet-4-6"),
        ("NLI verification",   "DeBERTa-v3 (MNLI/FEVER)",             "gpt-4o / claude-sonnet-4-6"),
        ("Output structuring", "Outlines / Pydantic",                  "n/a"),
    ],
)

doc.add_paragraph(
    "Serving: vLLM · FastAPI · Qdrant · Postgres · MinIO · Docker Compose. "
    "Orchestration: Prefect 2."
)

add_figure("fig03_crawl_console.png",       FIGURES[2][1])
add_figure("fig04_harvest_review.png",      FIGURES[3][1])
add_figure("fig05_extraction_ocr_diff.png", FIGURES[4][1])
add_figure("fig06_mapping_run.png",         FIGURES[5][1])

# ---- §3 ACCURACY, TRANSPARENCY, AND COST-EFFICIENCY --------------------
doc.add_heading("3. Accuracy, Transparency, and Cost-Efficiency", level=1)

doc.add_heading("Accuracy — CVR Loop", level=2)

add_figure(DIAGRAM_CVR[0], DIAGRAM_CVR[1], width_in=6.4)

doc.add_paragraph(
    "Every claim passes three sequential gates. G1 Span Match: verbatim span present "
    "character-for-character in source (length-proportional fuzzy tolerance in OCR "
    "regions only; edit distance logged). G2 NLI Entailment: DeBERTa-v3 confirms "
    "the span entails the classification above a gold-set-calibrated threshold. "
    "G3 Structural Plausibility: cited section exists, rubric operative predicates "
    "present, clause role is operative not definitional. Any gate failure rejects or "
    "routes to human review. Autonomy (L0–L3) controls human gating; CVR gates "
    "always run."
)

add_figure("fig07_source_trace.png", FIGURES[6][1])

doc.add_paragraph(
    "A hand-labeled gold set (~20 clauses per jurisdiction, BD / TH / SG) benchmarks "
    "classification F1 (target ≥ 0.75) and citation accuracy (target ≥ 0.95). The "
    "same set runs against a naive frontier-model baseline to quantify the value of "
    "verification. CVR rejection rate and inter-annotator agreement (Cohen's κ) "
    "are also reported."
)

doc.add_heading("Transparency — Provenance Bundle", level=2)
doc.add_paragraph(
    "Every record carries: source URL, SHA-256, retrieval timestamp, section, page, "
    "char offsets, bbox, verbatim quote, and model identity. Any third party can "
    "re-verify any claim by fetching the source, recomputing the hash, and re-running "
    "NLI locally from the provenance bundle."
)

add_figure("fig08_export_output.png", FIGURES[7][1])

doc.add_heading("Cost-Efficiency", level=2)
doc.add_paragraph(
    "Self-hosted default incurs zero per-token cost. VLM repair runs only on "
    "low-confidence OCR regions. Cloud escalation is opt-in, per-task, and budget-"
    "capped; on cap the pipeline falls back to local and logs the event. Per-run "
    "cost is reported by task and provider."
)

# ---- SAVE ----------------------------------------------------------------
doc.save(OUT_DOCX)
print(f"\nWrote {OUT_DOCX}")
